from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from zipfile import ZipFile

src = Path("AA_Ejemplo_ANALISIS_DE_DATOS.docx")
doc = Document(src)

print(f"PARAGRAPHS={len(doc.paragraphs)} TABLES={len(doc.tables)} INLINE_SHAPES={len(doc.inline_shapes)}")
print("\n=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    text = " ".join(p.text.split())
    if text:
        print(f"P{i:03d} [{p.style.name}] {text}")

print("\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print(f"TABLE {ti}: {len(table.rows)}x{len(table.columns)}")
    for ri, row in enumerate(table.rows):
        vals = [" ".join(cell.text.split()) for cell in row.cells]
        print(f"R{ri:02d}: " + " || ".join(vals))

media = Path("analysis_example_media")
media.mkdir(exist_ok=True)
with ZipFile(src) as z:
    names = [n for n in z.namelist() if n.startswith("word/media/")]
    print("\n=== MEDIA ===")
    for n in names:
        target = media / Path(n).name
        target.write_bytes(z.read(n))
        print(f"{target} {target.stat().st_size}")
